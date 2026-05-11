from __future__ import annotations

import os
import shutil
import threading
import time
import unittest
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from src.artifacts import (
    load_acquisition_result,
    load_input_source,
    persist_gpu_diagnostics,
    persist_source_artifacts,
    persist_transcript_artifacts,
)
from src.config import load_settings
from src.diagnostics import (
    check_api_config,
    check_python_version,
    check_whisper_cpp_portability,
    check_workspace_runtime_mode,
    overall_status,
)
from src.models import (
    AcquisitionResult,
    FirstPassResult,
    InputSource,
    LowValueSegment,
    TemplateDefinition,
    TemplateResult,
    TemplateSection,
    TranscriptBundle,
    TranscriptChunk,
    TranscriptSegment,
)
from src.progress import build_status, build_summary_input_preview
from src.result_writers import render_first_pass_markdown, render_template_markdown
from src.segmenter import build_chunks
from src.source_ingest import (
    acquire_text_file,
    classify_media_content_type,
    create_local_media_source,
    create_text_file_source,
)
from src.transcriber import (
    _build_gpu_diagnostic_entry,
    _capture_gpu_snapshot,
    _ensure_local_runtime_supported,
    _ensure_whisper_cpp_binary_portable,
    _run_local_worker,
    _should_chunk_local_audio,
    _wait_for_gpu_cleanup,
    _wait_for_worker_result_file,
)
from src.transcribers.whisper_cpp import WhisperCppTranscriber
from src.utils import hidden_subprocess_kwargs, run_command


class SegmenterTests(unittest.TestCase):
    def test_build_chunks_splits_on_gap(self) -> None:
        segments = [
            TranscriptSegment(index=1, start=0, end=10, text="part one"),
            TranscriptSegment(index=2, start=12, end=20, text="part two"),
            TranscriptSegment(index=3, start=50, end=60, text="part three"),
        ]
        chunks = build_chunks(segments, max_minutes=3, gap_seconds=12)
        self.assertEqual(len(chunks), 2)
        self.assertEqual(chunks[0].text, "part one\n\npart two")
        self.assertEqual(chunks[1].text, "part three")

    def test_build_chunks_without_timeline_uses_paragraph_grouping(self) -> None:
        segments = [
            TranscriptSegment(index=1, start=None, end=None, text="paragraph one"),
            TranscriptSegment(index=2, start=None, end=None, text="paragraph two"),
        ]
        chunks = build_chunks(segments, max_minutes=3, gap_seconds=12, max_chars=100)
        self.assertEqual(len(chunks), 1)
        self.assertIsNone(chunks[0].start)
        self.assertEqual(chunks[0].text, "paragraph one\n\nparagraph two")


class ResultWriterTests(unittest.TestCase):
    def test_render_first_pass_markdown_uses_source_metadata(self) -> None:
        source = InputSource(
            source_id="txt",
            entry_type="text_file",
            content_type="plain_text",
            display_name="demo.txt",
            original_filename="demo.txt",
        )
        acquisition = AcquisitionResult(source=source, acquisition_mode="direct_text", display_text="raw text")
        transcript_bundle = TranscriptBundle(
            source=source,
            acquisition=acquisition,
            output_dir=Path("output/demo"),
            backend_used="direct_text",
            model_used="text-input",
            language="zh",
            transcript_text="raw text",
        )
        first_pass = FirstPassResult(
            provider="fallback",
            model="rule-based",
            cleaned_transcript="cleaned text",
            one_line_verdict="值得看，核心价值在于 point a",
            headline_verdict="worth reading",
            value_rating="partial value",
            value_reason="some useful points",
            high_value_points=["point a"],
            objective_context=["context a"],
            low_value_segments=[],
            raw_transcript_reference="raw text",
        )

        rendered = render_first_pass_markdown(first_pass, transcript_bundle)
        self.assertIn("text_file", rendered)
        self.assertIn("direct_text", rendered)
        self.assertIn("demo.txt", rendered)

    def test_render_template_markdown_includes_template_fields(self) -> None:
        source = InputSource(
            source_id="txt",
            entry_type="pasted_text",
            content_type="plain_text",
            display_name="pasted_text",
        )
        acquisition = AcquisitionResult(source=source, acquisition_mode="direct_text", display_text="raw text")
        transcript_bundle = TranscriptBundle(
            source=source,
            acquisition=acquisition,
            output_dir=Path("output/demo"),
            backend_used="direct_text",
            model_used="text-input",
            language="zh",
            transcript_text="raw text",
        )
        first_pass = FirstPassResult(
            provider="fallback",
            model="rule-based",
            cleaned_transcript="cleaned text",
            one_line_verdict="值得看，核心价值在于 point a",
            headline_verdict="worth reading",
            value_rating="partial value",
            value_reason="some useful points",
            high_value_points=["point a"],
            objective_context=["context a"],
            low_value_segments=[LowValueSegment(start=None, end=None, reason="filler", excerpt="small talk")],
            raw_transcript_reference="raw text",
        )
        template_result = TemplateResult(
            template_id="course-notes",
            template_name="Course Notes",
            provider="fallback",
            model="rule-based",
            overview="overview",
            key_points=["point a"],
            section_summaries=[TemplateSection(title="section a", summary="summary a", bullets=["bullet a"])],
            template_fields={"review_notes": ["note a"]},
        )
        template_definition = TemplateDefinition(
            id="course-notes",
            name="Course Notes",
            description="notes",
            input_fields=[],
            output_fields=[],
            prompt_instructions="",
        )

        rendered = render_template_markdown(
            template_result,
            first_pass,
            transcript_bundle,
            template_definition,
        )
        self.assertIn("review_notes", rendered)
        self.assertIn("section a", rendered)


class ConfigTests(unittest.TestCase):
    def test_load_settings_reads_env_file(self) -> None:
        original_values = {
            "TRANSCRIBE_BACKEND": os.environ.get("TRANSCRIBE_BACKEND"),
            "SUMMARY_API_MODEL": os.environ.get("SUMMARY_API_MODEL"),
            "CHUNK_GAP_SECONDS": os.environ.get("CHUNK_GAP_SECONDS"),
        }
        for key in original_values:
            os.environ.pop(key, None)

        temp_dir = Path("output") / "test_env"
        temp_dir.mkdir(parents=True, exist_ok=True)
        try:
            env_path = temp_dir / ".env"
            env_path.write_text(
                "TRANSCRIBE_BACKEND=auto\nSUMMARY_API_MODEL=gpt-5.1\nCHUNK_GAP_SECONDS=8\n",
                encoding="utf-8",
            )
            settings = load_settings(env_path)
            self.assertEqual(settings.transcribe_backend, "auto")
            self.assertEqual(settings.summary_api_model, "gpt-5.1")
            self.assertEqual(settings.chunk_gap_seconds, 8)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
            for key, value in original_values.items():
                if value is None:
                    os.environ.pop(key, None)
                else:
                    os.environ[key] = value


class SourceIngestTests(unittest.TestCase):
    def test_classify_media_content_type(self) -> None:
        self.assertEqual(classify_media_content_type("demo.mp4"), "video")
        self.assertEqual(classify_media_content_type("demo.mp3"), "audio")
        self.assertIsNone(classify_media_content_type("demo.txt"))

    def test_create_local_media_source_rejects_text_file(self) -> None:
        run_dir = Path("output") / "bad_media_source_test"
        shutil.rmtree(run_dir, ignore_errors=True)
        run_dir.mkdir(parents=True, exist_ok=True)
        try:
            with self.assertRaises(RuntimeError):
                create_local_media_source(file_name="demo.txt", file_bytes=b"hello", run_dir=run_dir)
        finally:
            shutil.rmtree(run_dir, ignore_errors=True)

    def test_acquire_text_file_supports_plain_text(self) -> None:
        run_dir = Path("output") / "text_ingest_test"
        shutil.rmtree(run_dir, ignore_errors=True)
        run_dir.mkdir(parents=True, exist_ok=True)
        try:
            source, _ = create_text_file_source(
                file_name="demo.txt",
                file_bytes=b"first paragraph\n\nsecond paragraph",
                run_dir=run_dir,
            )
            acquisition = acquire_text_file(source)
            self.assertEqual(acquisition.acquisition_mode, "direct_text")
            self.assertEqual(len(acquisition.segments), 2)
            self.assertIsNone(acquisition.segments[0].start)
        finally:
            shutil.rmtree(run_dir, ignore_errors=True)

    def test_acquire_text_file_supports_subtitle_timeline(self) -> None:
        run_dir = Path("output") / "subtitle_ingest_test"
        shutil.rmtree(run_dir, ignore_errors=True)
        run_dir.mkdir(parents=True, exist_ok=True)
        try:
            source, _ = create_text_file_source(
                file_name="demo.vtt",
                file_bytes=b"WEBVTT\n\n00:00:00.000 --> 00:00:01.000\nhello\n",
                run_dir=run_dir,
            )
            acquisition = acquire_text_file(source)
            self.assertEqual(acquisition.acquisition_mode, "subtitle_parse")
            self.assertEqual(acquisition.display_text, "hello")
            self.assertEqual(acquisition.segments[0].start, 0.0)
            self.assertEqual(acquisition.segments[0].end, 1.0)
        finally:
            shutil.rmtree(run_dir, ignore_errors=True)

    def test_acquire_text_file_supports_docx(self) -> None:
        from docx import Document

        run_dir = Path("output") / "docx_ingest_test"
        shutil.rmtree(run_dir, ignore_errors=True)
        run_dir.mkdir(parents=True, exist_ok=True)
        try:
            document = Document()
            document.add_paragraph("Title")
            document.add_paragraph("Body one")
            document.add_paragraph("Body two")
            buffer = BytesIO()
            document.save(buffer)

            source, _ = create_text_file_source(
                file_name="demo.docx",
                file_bytes=buffer.getvalue(),
                run_dir=run_dir,
            )
            acquisition = acquire_text_file(source)
            self.assertEqual(acquisition.acquisition_mode, "document_text_extract")
            self.assertEqual(len(acquisition.segments), 3)
            self.assertIn("Body one", acquisition.display_text)
        finally:
            shutil.rmtree(run_dir, ignore_errors=True)


class ArtifactTests(unittest.TestCase):
    def test_persist_source_artifacts_writes_and_loads_files(self) -> None:
        run_dir = Path("output") / "artifact_source_test"
        shutil.rmtree(run_dir, ignore_errors=True)
        run_dir.mkdir(parents=True, exist_ok=True)
        try:
            source = InputSource(
                source_id="txt",
                entry_type="text_file",
                content_type="plain_text",
                display_name="demo.txt",
                original_filename="demo.txt",
            )
            acquisition = AcquisitionResult(
                source=source,
                acquisition_mode="direct_text",
                display_text="raw text",
                segments=[TranscriptSegment(index=1, start=None, end=None, text="raw text")],
            )
            bundle = TranscriptBundle(
                source=source,
                acquisition=acquisition,
                output_dir=run_dir,
                backend_used="direct_text",
                model_used="text-input",
                language="zh",
                transcript_text="raw text",
            )
            input_source_path, acquisition_result_path = persist_source_artifacts(run_dir, bundle)
            loaded_source = load_input_source(input_source_path)
            loaded_acquisition = load_acquisition_result(acquisition_result_path, loaded_source)
            self.assertEqual(loaded_source.entry_type, "text_file")
            self.assertEqual(loaded_acquisition.acquisition_mode, "direct_text")
        finally:
            shutil.rmtree(run_dir, ignore_errors=True)

    def test_persist_transcript_artifacts_writes_files_before_summary(self) -> None:
        run_dir = Path("output") / "artifact_test"
        run_dir.mkdir(parents=True, exist_ok=True)
        try:
            source = InputSource(
                source_id="txt",
                entry_type="text_file",
                content_type="plain_text",
                display_name="demo.txt",
            )
            acquisition = AcquisitionResult(source=source, acquisition_mode="direct_text", display_text="bundle text")
            transcript_bundle = TranscriptBundle(
                source=source,
                acquisition=acquisition,
                output_dir=run_dir,
                backend_used="direct_text",
                model_used="text-input",
                language="zh",
                transcript_text="bundle text",
                segments=[TranscriptSegment(index=1, start=None, end=None, text="bundle text")],
                chunks=[TranscriptChunk(index=1, start=None, end=None, text="bundle text")],
            )
            transcript_path, segments_path, chunks_path = persist_transcript_artifacts(run_dir, transcript_bundle)
            self.assertTrue(transcript_path.exists())
            self.assertTrue(segments_path.exists())
            self.assertTrue(chunks_path.exists())
        finally:
            shutil.rmtree(run_dir, ignore_errors=True)

    def test_persist_gpu_diagnostics_writes_files(self) -> None:
        run_dir = Path("output") / "gpu_diag_test"
        run_dir.mkdir(parents=True, exist_ok=True)
        try:
            json_path, md_path = persist_gpu_diagnostics(
                run_dir,
                [
                    {
                        "timestamp": "2026-04-24 10:00:00",
                        "event": "chunk_start",
                        "chunk_index": 1,
                        "chunk_total": 5,
                        "device": "cuda",
                        "worker_pid": 123,
                        "return_code": None,
                        "gpu_memory_used_mb": 1024,
                        "gpu_utilization_percent": 50,
                        "result_exists": None,
                        "segment_count": None,
                        "note": "start",
                    }
                ],
            )
            self.assertTrue(json_path.exists())
            self.assertTrue(md_path.exists())
        finally:
            shutil.rmtree(run_dir, ignore_errors=True)


class DiagnosticsTests(unittest.TestCase):
    def test_hidden_subprocess_kwargs_include_windows_flags(self) -> None:
        with patch("src.utils.os.name", "nt"):
            kwargs = hidden_subprocess_kwargs()
        self.assertIn("creationflags", kwargs)
        self.assertIn("startupinfo", kwargs)

    def test_run_command_uses_hidden_subprocess_spawn(self) -> None:
        fake_result = SimpleNamespace(stdout="", stderr="", returncode=0)
        with (
            patch("src.utils.hidden_subprocess_kwargs", return_value={"creationflags": 99}),
            patch("src.utils.subprocess.run", return_value=fake_result) as mock_run,
        ):
            result = run_command(["ffmpeg", "-version"], "ffmpeg failed")

        self.assertIs(result, fake_result)
        self.assertEqual(mock_run.call_args.kwargs["creationflags"], 99)

    def test_capture_gpu_snapshot_uses_hidden_subprocess_spawn(self) -> None:
        fake_completed = SimpleNamespace(stdout="1024, 50\n")
        with (
            patch("src.transcriber.shutil.which", return_value="nvidia-smi.exe"),
            patch("src.transcriber.hidden_subprocess_kwargs", return_value={"creationflags": 88}),
            patch("src.transcriber.subprocess.run", return_value=fake_completed) as mock_run,
        ):
            snapshot = _capture_gpu_snapshot()

        self.assertEqual(snapshot["gpu_memory_used_mb"], 1024)
        self.assertEqual(snapshot["gpu_utilization_percent"], 50)
        self.assertEqual(mock_run.call_args.kwargs["creationflags"], 88)

    def test_local_worker_uses_hidden_subprocess_spawn(self) -> None:
        temp_dir = Path("output") / "worker_hidden_spawn_test"
        temp_dir.mkdir(parents=True, exist_ok=True)
        audio_path = temp_dir / "audio.wav"
        audio_path.write_bytes(b"RIFF")

        class FakeHandle:
            pid = 1234
            returncode = 0

            def poll(self) -> int:
                return 0

            def wait(self) -> int:
                return 0

        settings = load_settings()
        settings.whisper_cpp_bin = str((temp_dir / "whisper-cli.exe").resolve())
        Path(settings.whisper_cpp_bin).write_text("stub", encoding="utf-8")
        settings.whisper_cpp_model_path = temp_dir / "model.bin"
        settings.whisper_cpp_model_path.write_text("stub", encoding="utf-8")
        settings.local_whisper_language = "zh"
        settings.local_whisper_prompt = ""

        try:
            with (
                patch("src.transcriber.hidden_subprocess_kwargs", return_value={"creationflags": 77}),
                patch("src.transcriber.subprocess.Popen", return_value=FakeHandle()) as mock_popen,
                patch("src.transcriber._capture_gpu_snapshot", return_value={"gpu_memory_used_mb": None, "gpu_utilization_percent": None}),
                patch("src.transcriber._get_wav_duration", return_value=1.0),
            ):
                result_path = temp_dir / "audio_whispercpp.json"
                result_path.write_text(
                    '{"language":"zh","transcription":[{"offsets":{"from":0,"to":1000},"text":"hello"}]}',
                    encoding="utf-8",
                )
                result = _run_local_worker(
                    audio_path=audio_path,
                    settings=settings,
                    total_duration=1.0,
                    progress_hook=None,
                )
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

        self.assertEqual(result["return_code"], 0)
        self.assertEqual(mock_popen.call_args.kwargs["creationflags"], 77)

    def test_check_python_version_warns_for_newer_runtime(self) -> None:
        self.assertEqual(check_python_version((3, 14, 0)).status, "warn")

    def test_check_api_config_warns_when_missing_key(self) -> None:
        item = check_api_config(
            name="Summary API",
            api_key=None,
            model="gpt-5.4",
            base_url=None,
            backend_note="summary phase uses a remote model",
        )
        self.assertEqual(item.status, "warn")

    def test_check_whisper_cpp_portability_flags_debug_runtime(self) -> None:
        temp_dir = Path("output") / "whisper_portability_test"
        temp_dir.mkdir(parents=True, exist_ok=True)
        binary_path = temp_dir / "whisper-cli.exe"
        binary_path.write_bytes(b"stub...MSVCP140D.dll...ucrtbased.dll")
        settings = load_settings()
        settings.whisper_cpp_bin = str(binary_path)

        try:
            item = check_whisper_cpp_portability(settings)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

        self.assertEqual(item.status, "error")
        self.assertIn("MSVCP140D.dll", item.detail)

    def test_workspace_runtime_check_reports_fastapi_workspace(self) -> None:
        item = check_workspace_runtime_mode()
        self.assertEqual(item.status, "ok")
        self.assertIn("FastAPI + workspace-ui", item.detail)

    def test_overall_status_prefers_error(self) -> None:
        status = overall_status(
            [
                check_python_version((3, 11, 9)),
                check_api_config(
                    name="Summary API",
                    api_key=None,
                    model="gpt-5.4",
                    base_url=None,
                    backend_note="summary phase uses a remote model",
                ),
                check_python_version((3, 14, 0)),
            ]
        )
        self.assertEqual(status, "warn")

    def test_local_runtime_guard_accepts_python_314(self) -> None:
        _ensure_local_runtime_supported((3, 14, 3))

    def test_whisper_cpp_portability_guard_rejects_debug_runtime(self) -> None:
        temp_dir = Path("output") / "whisper_portability_guard_test"
        temp_dir.mkdir(parents=True, exist_ok=True)
        binary_path = temp_dir / "whisper-cli.exe"
        binary_path.write_bytes(b"stub...VCRUNTIME140D.dll")

        try:
            with self.assertRaises(RuntimeError) as context:
                _ensure_whisper_cpp_binary_portable(str(binary_path))
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

        self.assertIn("Debug CRT", str(context.exception))

    def test_whisper_cpp_parser_reads_segment_offsets(self) -> None:
        temp_dir = Path("output") / "whisper_cpp_parse_test"
        temp_dir.mkdir(parents=True, exist_ok=True)
        json_path = temp_dir / "sample.json"
        try:
            json_path.write_text(
                """
                {
                  "language": "zh",
                  "transcription": [
                    {
                      "offsets": { "from": 0, "to": 380 },
                      "text": "hello"
                    },
                    {
                      "offsets": { "from": 380, "to": 3000 },
                      "text": "test"
                    }
                  ]
                }
                """.strip(),
                encoding="utf-8",
            )
            transcriber = WhisperCppTranscriber(
                binary_path="whisper-cli",
                model_path=Path("models/ggml-medium.bin"),
                language="zh",
                prompt="",
            )
            segments, language, notes = transcriber.load_result(json_path=json_path)
            self.assertEqual(language, "zh")
            self.assertEqual(notes, [])
            self.assertEqual(len(segments), 2)
            self.assertEqual(segments[0].text, "hello")
            self.assertAlmostEqual(segments[0].end, 0.38)
            self.assertAlmostEqual(segments[1].start, 0.38)
            self.assertAlmostEqual(segments[1].end, 3.0)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def test_whisper_cpp_build_command_uses_absolute_paths(self) -> None:
        transcriber = WhisperCppTranscriber(
            binary_path="whisper-cli",
            model_path=Path("models/ggml-medium.bin"),
            language="zh",
            prompt="",
        )
        command = transcriber.build_command(
            audio_path=Path("output/demo/audio.wav"),
            output_prefix=Path("output/demo/audio_whispercpp"),
        )
        self.assertIn(str(Path("models/ggml-medium.bin").resolve()), command)
        self.assertIn(str(Path("output/demo/audio.wav").resolve()), command)
        self.assertIn(str(Path("output/demo/audio_whispercpp").resolve()), command)

    def test_whisper_cpp_parser_repairs_latin1_wrapped_utf8_text(self) -> None:
        temp_dir = Path("output") / "whisper_cpp_parse_latin1_test"
        temp_dir.mkdir(parents=True, exist_ok=True)
        json_path = temp_dir / "sample.json"
        try:
            json_path.write_bytes(
                (
                    b'{"language":"zh","transcription":[{"offsets":{"from":0,"to":1000},'
                    b'"text":"'
                    + "hello".encode("utf-8")
                    + b'","tokens":[{"text":"\xe5\x86"}]}]}'
                )
            )
            transcriber = WhisperCppTranscriber(
                binary_path="whisper-cli",
                model_path=Path("models/ggml-medium.bin"),
                language="zh",
                prompt="",
            )
            segments, language, _ = transcriber.load_result(json_path=json_path)
            self.assertEqual(language, "zh")
            self.assertEqual(len(segments), 1)
            self.assertEqual(segments[0].text, "hello")
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def test_local_gpu_chunking_enabled_for_long_audio(self) -> None:
        self.assertTrue(_should_chunk_local_audio(total_duration=1800, device="cuda"))
        self.assertFalse(_should_chunk_local_audio(total_duration=300, device="cuda"))
        self.assertTrue(_should_chunk_local_audio(total_duration=1800, device="cpu"))

    def test_wait_for_worker_result_file_handles_delayed_write(self) -> None:
        temp_dir = Path("output") / "worker_result_wait_test"
        temp_dir.mkdir(parents=True, exist_ok=True)
        result_path = temp_dir / "result.json"
        try:

            def delayed_write() -> None:
                time.sleep(0.2)
                result_path.write_text('{"ok": true}', encoding="utf-8")

            writer = threading.Thread(target=delayed_write)
            writer.start()
            payload = _wait_for_worker_result_file(result_path, timeout_seconds=1.0, poll_interval_seconds=0.05)
            writer.join()
            self.assertEqual(payload, {"ok": True})
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def test_wait_for_worker_result_file_accepts_latin1_json(self) -> None:
        temp_dir = Path("output") / "worker_result_latin1_test"
        temp_dir.mkdir(parents=True, exist_ok=True)
        result_path = temp_dir / "result.json"
        try:
            result_path.write_bytes(
                (
                    b'{"language":"zh","transcription":[{"offsets":{"from":0,"to":1000},"text":"'
                    + "hello".encode("utf-8")
                    + b'","tokens":[{"text":"\xe5\x86"}]}]}'
                )
            )
            payload = _wait_for_worker_result_file(result_path, timeout_seconds=0.1, poll_interval_seconds=0.01)
            self.assertIsInstance(payload, dict)
            self.assertEqual(payload["language"], "zh")
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def test_capture_gpu_snapshot_returns_known_shape(self) -> None:
        snapshot = _capture_gpu_snapshot()
        self.assertIn("gpu_memory_used_mb", snapshot)
        self.assertIn("gpu_utilization_percent", snapshot)

    def test_wait_for_gpu_cleanup_returns_diagnostic_entry(self) -> None:
        settings = load_settings()
        settings.gpu_cleanup_wait_seconds = 0
        entry = _wait_for_gpu_cleanup(settings)
        self.assertIn(entry["event"], {"gpu_cleanup_observed", "gpu_cleanup_timeout"})
        self.assertEqual(entry["device"], "cuda")

    def test_build_gpu_diagnostic_entry_contains_expected_fields(self) -> None:
        entry = _build_gpu_diagnostic_entry(
            event="chunk_start",
            chunk_index=1,
            chunk_total=5,
            device="cuda",
            worker_pid=123,
            return_code=0,
            gpu_snapshot={"gpu_memory_used_mb": 1024, "gpu_utilization_percent": 50},
            result_exists=True,
            segment_count=10,
            note="ok",
        )
        self.assertEqual(entry["chunk_index"], 1)
        self.assertEqual(entry["worker_pid"], 123)
        self.assertEqual(entry["gpu_memory_used_mb"], 1024)


class ProgressTests(unittest.TestCase):
    def test_build_status_sequence_progresses(self) -> None:
        created = build_status("task_created", "created")
        transcribing = build_status("transcribing", "running")
        completed = build_status("completed", "done")
        self.assertLess(created.progress_percent, transcribing.progress_percent)
        self.assertLess(transcribing.progress_percent, completed.progress_percent)

    def test_build_status_keeps_gpu_metadata(self) -> None:
        status = build_status(
            "transcribing",
            "still running",
            worker_pid=123,
            gpu_memory_used_mb=2048,
            last_heartbeat_at="2026-04-24 10:00:00",
            last_segment_at="2026-04-24 09:59:00",
            active_chunk_index=5,
            active_chunk_total=5,
        )
        self.assertEqual(status.worker_pid, 123)
        self.assertEqual(status.gpu_memory_used_mb, 2048)
        self.assertEqual(status.active_chunk_index, 5)

    def test_summary_input_preview_uses_chunk_text(self) -> None:
        source = InputSource(
            source_id="txt",
            entry_type="text_file",
            content_type="plain_text",
            display_name="demo.txt",
        )
        acquisition = AcquisitionResult(source=source, acquisition_mode="direct_text", display_text="overall text")
        transcript_bundle = TranscriptBundle(
            source=source,
            acquisition=acquisition,
            output_dir=Path("output/demo"),
            backend_used="direct_text",
            model_used="text-input",
            language="zh",
            transcript_text="overall text",
            chunks=[TranscriptChunk(index=1, start=None, end=None, text="preview text for summary model")],
        )
        preview = build_summary_input_preview(transcript_bundle)
        self.assertIn("preview text for summary model", preview)


if __name__ == "__main__":
    unittest.main()
